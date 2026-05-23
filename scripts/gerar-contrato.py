from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1, color=RGBColor(0x1a, 0x1a, 0x5e)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12) if level == 1 else Pt(11)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a1a5e')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─────────────────────────────────────────────────────────────────────────────
# CABEÇALHO
# ─────────────────────────────────────────────────────────────────────────────
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = titulo.add_run("CONTRATO DE LICENÇA E PRESTAÇÃO DE SERVIÇOS")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)

subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitulo.add_run("Sistema de Gestão para Salão de Beleza")
r2.bold = True
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
add_line()
doc.add_paragraph()

data_hoje = datetime.date.today().strftime("%d de %B de %Y").replace(
    "January","janeiro").replace("February","fevereiro").replace(
    "March","março").replace("April","abril").replace("May","maio").replace(
    "June","junho").replace("July","julho").replace("August","agosto").replace(
    "September","setembro").replace("October","outubro").replace(
    "November","novembro").replace("December","dezembro")

body(f"Pelo presente instrumento particular, as partes abaixo qualificadas celebram o presente "
     f"Contrato de Licença de Software e Prestação de Serviços ("Contrato"), que se regerá pelas "
     f"cláusulas e condições a seguir estipuladas.")

# ─────────────────────────────────────────────────────────────────────────────
# QUALIFICAÇÃO DAS PARTES
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 1ª – DAS PARTES")

body("CONTRATADO (Desenvolvedor):", bold=True)
rows_contratado = [
    ("Nome / Razão Social", "_______________________________________________"),
    ("CPF / CNPJ",          "_______________________________________________"),
    ("Endereço",            "_______________________________________________"),
    ("Cidade / Estado",     "_______________________________________________"),
    ("E-mail",              "_______________________________________________"),
    ("Telefone",            "_______________________________________________"),
]
t = doc.add_table(rows=len(rows_contratado), cols=2)
t.style = "Table Grid"
for i, (label, val) in enumerate(rows_contratado):
    t.rows[i].cells[0].text = label
    t.rows[i].cells[1].text = val
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(10.5)
doc.add_paragraph()

body("CONTRATANTE (Comprador):", bold=True)
rows_contratante = [
    ("Nome / Razão Social", "_______________________________________________"),
    ("CPF / CNPJ",          "_______________________________________________"),
    ("Endereço",            "_______________________________________________"),
    ("Cidade / Estado",     "_______________________________________________"),
    ("E-mail",              "_______________________________________________"),
    ("Telefone",            "_______________________________________________"),
    ("Nome do Salão",       "_______________________________________________"),
]
t2 = doc.add_table(rows=len(rows_contratante), cols=2)
t2.style = "Table Grid"
for i, (label, val) in enumerate(rows_contratante):
    t2.rows[i].cells[0].text = label
    t2.rows[i].cells[1].text = val
    t2.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t2.columns[0].width = Cm(5)
    t2.columns[1].width = Cm(10.5)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# OBJETO
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 2ª – DO OBJETO")
body("2.1. O presente Contrato tem por objeto a licença de uso do software denominado "
     "Sistema de Gestão para Salão de Beleza ("Software"), desenvolvido pelo CONTRATADO, "
     "composto pelas seguintes funcionalidades:")

funcionalidades = [
    "Página pública (landing page) com apresentação do salão, serviços e preços;",
    "Sistema de agendamento online para clientes com seleção de serviço, data e horário;",
    "Painel administrativo (dashboard) com gestão completa de agendamentos;",
    "Cadastro e gestão de serviços, categorias e tabela de preços;",
    "Gestão de usuários (clientes e administradores);",
    "Módulo de avaliações e comentários dos clientes;",
    "Controle de vendas e histórico financeiro;",
    "Bloqueio de horários e períodos no calendário;",
    "Notificações via WhatsApp (confirmação e cancelamento de agendamentos);",
    "Sistema de recuperação de senha por e-mail;",
    "Gerenciamento de banners, rodapé e configurações visuais do site;",
    "Armazenamento em banco de dados PostgreSQL com persistência de imagens em Base64.",
]
for f in funcionalidades:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
body("2.2. O CONTRATADO entregará o Software instalado, configurado e em pleno funcionamento "
     "no ambiente de hospedagem acordado entre as partes.")

# ─────────────────────────────────────────────────────────────────────────────
# VALOR E PAGAMENTO
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 3ª – DO VALOR E FORMA DE PAGAMENTO")
body("3.1. Pela licença de uso do Software e pelos serviços de instalação, configuração e "
     "treinamento inicial, o CONTRATANTE pagará ao CONTRATADO o valor total de:")

valor_p = doc.add_paragraph()
valor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
vr = valor_p.add_run("R$ _______________ (_________________________________ reais)")
vr.bold = True
vr.font.size = Pt(13)
vr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)

body("3.2. O pagamento será realizado da seguinte forma:")
body("(  ) À vista, mediante PIX / TED / boleto, até a data de entrega do Software.", indent=True)
body("(  ) Em _____ parcelas de R$ _______________, sendo a 1ª na entrega e as demais "
     "com vencimento no mesmo dia dos meses subsequentes.", indent=True)

body("3.3. Em caso de atraso no pagamento, incidirão multa de 2% (dois por cento) sobre o "
     "valor em aberto, acrescida de juros de 1% (um por cento) ao mês, pro rata die.")

body("3.4. Os dados para pagamento são:")
body("Banco: ___________  Agência: __________  Conta: __________  "
     "Tipo: __________  PIX: ________________________", indent=True)

# ─────────────────────────────────────────────────────────────────────────────
# ENTREGA
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 4ª – DA ENTREGA E INSTALAÇÃO")
body("4.1. O CONTRATADO se compromete a entregar o Software instalado e operacional no prazo "
     "de _____ (___________) dias úteis, contados da assinatura deste Contrato e do "
     "recebimento do primeiro pagamento.")
body("4.2. A entrega compreende: instalação em servidor, configuração inicial do sistema "
     "(cadastro do salão, serviços e horários de funcionamento) e treinamento remoto de até "
     "2 (duas) horas para o CONTRATANTE.")
body("4.3. Eventuais customizações não previstas neste Contrato serão orçadas e cobradas "
     "separadamente, mediante aprovação prévia do CONTRATANTE.")

# ─────────────────────────────────────────────────────────────────────────────
# SUPORTE
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 5ª – DO SUPORTE E MANUTENÇÃO")
body("5.1. O CONTRATADO oferece suporte técnico gratuito por _____ (___________) dias corridos "
     "a partir da data de entrega, para correção de eventuais falhas do Software.")
body("5.2. Após o período de suporte gratuito, a manutenção poderá ser contratada "
     "mediante mensalidade de R$ _______________ ou por demanda, conforme tabela vigente "
     "do CONTRATADO.")
body("5.3. O suporte será prestado por e-mail ou WhatsApp, em dias úteis, no horário das "
     "09h às 18h, com prazo de resposta de até 24 (vinte e quatro) horas.")
body("5.4. Não são cobertos pelo suporte gratuito: problemas decorrentes de uso indevido, "
     "alterações realizadas por terceiros no código-fonte, falhas de infraestrutura de "
     "hospedagem ou de conectividade com a internet.")

# ─────────────────────────────────────────────────────────────────────────────
# PROPRIEDADE INTELECTUAL
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 6ª – DA PROPRIEDADE INTELECTUAL E LICENÇA")
body("6.1. O Software é de propriedade intelectual exclusiva do CONTRATADO, protegido pela "
     "Lei nº 9.609/1998 (Lei do Software) e pela Lei nº 9.610/1998 (Lei de Direitos Autorais).")

body("6.2. O presente Contrato confere ao CONTRATANTE:")
itens_licenca = [
    "Licença de uso não exclusiva e intransferível do Software;",
    "Direito de operar o Software em ambiente de produção de um único estabelecimento;",
    "Acesso ao código-fonte, caso expressamente indicado no item 6.3 abaixo.",
]
for item in itens_licenca:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
body("6.3. Cessão de código-fonte:")
body("(  ) O código-fonte NÃO está incluído neste Contrato. O CONTRATANTE recebe apenas "
     "o Software instalado e em funcionamento.", indent=True)
body("(  ) O código-fonte está incluído, sendo cedido ao CONTRATANTE para uso próprio, "
     "vedada a revenda ou sublicença a terceiros.", indent=True)

body("6.4. É vedado ao CONTRATANTE sublicenciar, revender, alugar, transferir ou de qualquer "
     "forma ceder a terceiros os direitos ora outorgados, sem prévia autorização escrita "
     "do CONTRATADO.")

# ─────────────────────────────────────────────────────────────────────────────
# HOSPEDAGEM
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 7ª – DA HOSPEDAGEM E INFRAESTRUTURA")
body("7.1. Os custos de hospedagem (servidor, domínio, banco de dados em nuvem e serviços "
     "de terceiros, como WhatsApp Business API e SendGrid) são de responsabilidade exclusiva "
     "do CONTRATANTE, salvo se expressamente incluídos no valor desta contratação.")
body("7.2. O CONTRATADO orienta e auxilia na configuração da hospedagem, mas não responde "
     "por interrupções causadas por terceiros (provedores de nuvem, operadoras, etc.).")

# ─────────────────────────────────────────────────────────────────────────────
# CONFIDENCIALIDADE
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 8ª – DA CONFIDENCIALIDADE")
body("8.1. As partes se comprometem a manter sigilo sobre informações confidenciais trocadas "
     "em razão deste Contrato, incluindo dados de clientes, credenciais de acesso e "
     "detalhes técnicos do Software.")
body("8.2. A obrigação de sigilo persiste por 2 (dois) anos após o término deste Contrato.")

# ─────────────────────────────────────────────────────────────────────────────
# RESCISÃO
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 9ª – DA RESCISÃO")
body("9.1. O presente Contrato poderá ser rescindido:")
rescisao = [
    "Por inadimplemento de qualquer das partes, após notificação prévia de 10 (dez) dias corridos;",
    "Por acordo mútuo, formalizado por escrito;",
    "Por descumprimento de qualquer cláusula contratual.",
]
for r in rescisao:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(r).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
body("9.2. Em caso de desistência pelo CONTRATANTE após o início do desenvolvimento ou "
     "instalação, será devida ao CONTRATADO multa rescisória equivalente a 30% (trinta "
     "por cento) do valor total contratado, além dos valores já pagos.")
body("9.3. Em caso de descumprimento pelo CONTRATADO sem justificativa, este restituirá "
     "os valores recebidos, deduzidos os serviços efetivamente prestados.")

# ─────────────────────────────────────────────────────────────────────────────
# DISPOSIÇÕES GERAIS
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 10ª – DAS DISPOSIÇÕES GERAIS")
body("10.1. Este Contrato representa o acordo integral entre as partes, substituindo qualquer "
     "entendimento anterior, verbal ou escrito, sobre o mesmo objeto.")
body("10.2. Qualquer alteração a este Contrato somente será válida se feita por escrito e "
     "assinada por ambas as partes.")
body("10.3. A tolerância de uma das partes quanto ao descumprimento de qualquer obrigação "
     "pela outra não implica renúncia ao direito de exigir o cumprimento futuro.")
body("10.4. Caso qualquer disposição deste Contrato seja considerada inválida ou inexequível, "
     "as demais cláusulas permanecerão em pleno vigor.")

# ─────────────────────────────────────────────────────────────────────────────
# FORO
# ─────────────────────────────────────────────────────────────────────────────
heading("CLÁUSULA 11ª – DO FORO")
body("11.1. As partes elegem o Foro da Comarca de ___________________________, Estado de "
     "_______________, para dirimir quaisquer dúvidas ou litígios oriundos deste Contrato, "
     "renunciando a qualquer outro, por mais privilegiado que seja.")

# ─────────────────────────────────────────────────────────────────────────────
# ASSINATURAS
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
add_line()

local_data = doc.add_paragraph()
local_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
local_data.add_run(
    f"___________________________, ___ de _________________ de ________"
).font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# Tabela de assinaturas
ts = doc.add_table(rows=4, cols=2)
ts.alignment = WD_TABLE_ALIGNMENT.CENTER

nomes = ["CONTRATADO", "CONTRATANTE"]
for col_idx, nome in enumerate(nomes):
    ts.rows[0].cells[col_idx].text = "_" * 40
    ts.rows[1].cells[col_idx].text = nome
    ts.rows[2].cells[col_idx].text = "Nome: _______________________________"
    ts.rows[3].cells[col_idx].text = "CPF / CNPJ: _________________________"
    for row_idx in range(4):
        cell = ts.rows[row_idx].cells[col_idx]
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
                if row_idx == 1:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Testemunhas
tw = doc.add_table(rows=3, cols=2)
tw.alignment = WD_TABLE_ALIGNMENT.CENTER
tit = tw.rows[0].cells[0]
tit.merge(tw.rows[0].cells[1])
tit.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
tit.paragraphs[0].add_run("TESTEMUNHAS").bold = True

for col_idx in range(2):
    tw.rows[1].cells[col_idx].text = f"1ª Testemunha:" if col_idx == 0 else "2ª Testemunha:"
    tw.rows[2].cells[col_idx].text = "Nome: ___________________________\nCPF:  ___________________________"
    for row_idx in range(1, 3):
        cell = tw.rows[row_idx].cells[col_idx]
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)

# ─────────────────────────────────────────────────────────────────────────────
# RODAPÉ
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
rodape = doc.add_paragraph()
rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rod = rodape.add_run(
    "Este documento possui _____ (___________) páginas e foi gerado digitalmente.\n"
    "Contrato de Licença de Software — Sistema de Gestão para Salão de Beleza"
)
r_rod.font.size = Pt(9)
r_rod.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ─────────────────────────────────────────────────────────────────────────────
# SALVAR
# ─────────────────────────────────────────────────────────────────────────────
output_path = "contrato-sistema-salao.docx"
doc.save(output_path)
print(f"Contrato gerado: {output_path}")
